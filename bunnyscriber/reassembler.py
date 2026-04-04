"""
Phase 5: Reassembly of per-speaker transcripts into a unified document.

Merges all speaker transcripts ordered by timestamp, with speaker labels,
crosstalk notation, and markers for uncertain segments.
"""

import os
from dataclasses import dataclass, field
from typing import List, Dict, Optional

from bunnyscriber.transcriber import SpeakerTranscript
from bunnyscriber.backends.base import TranscriptSegment


@dataclass
class UnifiedSegment:
    """A segment in the unified transcript with speaker info."""

    speaker_name: str
    text: str
    start: float
    end: float
    is_crosstalk: bool = False
    is_uncertain: bool = False
    is_non_speaker: bool = False
    overlap_with: Optional[str] = None


def reassemble_transcripts(
    transcripts: List[SpeakerTranscript],
    uncertain_labels: Optional[Dict[int, str]] = None,
    non_speaker_labels: Optional[set] = None,
    chunk_offsets: Optional[Dict[int, float]] = None,
) -> List[UnifiedSegment]:
    """Merge per-speaker transcripts into a single timeline.

    Args:
        transcripts: All speaker transcripts across all chunks.
        uncertain_labels: Dict mapping uncertain segment index to assigned speaker
                         or None if skipped.
        non_speaker_labels: Set of speaker labels flagged as non-speaker audio.
        chunk_offsets: Dict of chunk_index -> time offset in seconds from file start.

    Returns:
        List of UnifiedSegment ordered by timestamp.
    """
    if uncertain_labels is None:
        uncertain_labels = {}
    if non_speaker_labels is None:
        non_speaker_labels = set()
    if chunk_offsets is None:
        chunk_offsets = {}

    all_segments: List[UnifiedSegment] = []

    for transcript in transcripts:
        offset = chunk_offsets.get(transcript.chunk_index, 0.0)
        is_non_speaker = transcript.speaker_label in non_speaker_labels

        for seg in transcript.result.segments:
            all_segments.append(UnifiedSegment(
                speaker_name=transcript.speaker_name,
                text=seg.text,
                start=seg.start + offset,
                end=seg.end + offset,
                is_non_speaker=is_non_speaker,
            ))

    # Sort by start time
    all_segments.sort(key=lambda s: s.start)

    # Detect crosstalk using a sliding window of active segments.
    # Since segments are sorted by start time, we maintain a window of
    # segments that haven't ended yet and only compare within that window.
    active = []  # list of indices into all_segments
    for i, seg in enumerate(all_segments):
        # Remove segments from the window that have ended before this one starts
        active = [j for j in active if all_segments[j].end > seg.start]

        # Check for overlap with any active segment from a different speaker
        for j in active:
            seg_a = all_segments[j]
            if seg_a.speaker_name != seg.speaker_name:
                seg_a.is_crosstalk = True
                seg_a.overlap_with = seg.speaker_name
                seg.is_crosstalk = True
                seg.overlap_with = seg_a.speaker_name

        active.append(i)

    return all_segments


def format_timestamp(seconds: float) -> str:
    """Format seconds as HH:MM:SS."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def write_txt(
    segments: List[UnifiedSegment],
    output_path: str,
    title: str = "BunnyScriber Transcript",
) -> str:
    """Write the unified transcript as a plain text file.

    Args:
        segments: Ordered list of unified segments.
        output_path: Path for the output .txt file.
        title: Title to put at the top of the transcript.

    Returns:
        Path to the written file.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write(f"{'=' * len(title)}\n\n")

        current_speaker = None

        for seg in segments:
            ts = format_timestamp(seg.start)

            if seg.is_non_speaker:
                f.write(f"[{ts}] [AUDIO CLIP] {seg.text}\n\n")
                current_speaker = None
                continue

            prefix = ""
            suffix = ""

            if seg.is_crosstalk and seg.overlap_with:
                suffix = f" [CROSSTALK with {seg.overlap_with}]"

            if seg.is_uncertain:
                prefix = "[UNIDENTIFIED] "

            if seg.speaker_name != current_speaker:
                f.write(f"\n[{ts}] {prefix}{seg.speaker_name}:{suffix}\n")
                current_speaker = seg.speaker_name

            f.write(f"  {seg.text}{suffix if current_speaker == seg.speaker_name and seg.is_crosstalk else ''}\n")

    return output_path


def write_srt(
    segments: List[UnifiedSegment],
    output_path: str,
) -> str:
    """Write the unified transcript as an SRT subtitle file.

    Args:
        segments: Ordered list of unified segments.
        output_path: Path for the output .srt file.

    Returns:
        Path to the written file.
    """
    def _srt_time(seconds: float) -> str:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        return f"{hours:02d}:{minutes:02d}:{secs:06.3f}".replace(".", ",")

    with open(output_path, "w", encoding="utf-8") as f:
        for i, seg in enumerate(segments, 1):
            start = _srt_time(seg.start)
            end = _srt_time(seg.end)
            speaker = seg.speaker_name
            if seg.is_non_speaker:
                speaker = "[AUDIO CLIP]"
            elif seg.is_uncertain:
                speaker = "[UNIDENTIFIED]"

            f.write(f"{i}\n")
            f.write(f"{start} --> {end}\n")
            f.write(f"[{speaker}] {seg.text}\n\n")

    return output_path


def write_docx(
    segments: List[UnifiedSegment],
    output_path: str,
    title: str = "BunnyScriber Transcript",
) -> str:
    """Write the unified transcript as a DOCX file.

    Args:
        segments: Ordered list of unified segments.
        output_path: Path for the output .docx file.
        title: Document title.

    Returns:
        Path to the written file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_speaker = None

    for seg in segments:
        ts = format_timestamp(seg.start)

        if seg.is_non_speaker:
            p = doc.add_paragraph()
            run = p.add_run(f"[{ts}] [AUDIO CLIP] ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
            p.add_run(seg.text)
            current_speaker = None
            continue

        if seg.speaker_name != current_speaker:
            p = doc.add_paragraph()
            label = seg.speaker_name
            if seg.is_uncertain:
                label = f"[UNIDENTIFIED]"
            run = p.add_run(f"[{ts}] {label}:")
            run.bold = True
            run.font.color.rgb = RGBColor(0x83, 0x18, 0x43)
            current_speaker = seg.speaker_name

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        text = seg.text
        if seg.is_crosstalk and seg.overlap_with:
            text += f" [CROSSTALK with {seg.overlap_with}]"
        p.add_run(text)

    doc.save(output_path)
    return output_path
